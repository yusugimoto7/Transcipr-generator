"""Put invisible marker tokens into the TR and PR retainer Word files.

Each token marks where a Sign field goes. Tokens are rendered in white so
the PDF looks blank there, but pdfplumber can still read their position.
"""
import copy
import docx
from docx.shared import RGBColor, Pt

WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def tok(name):
    return f"ZQ{name}ZQ"


def _style_token(run):
    run.font.color.rgb = WHITE
    run.font.size = Pt(7)
    run.font.bold = False


def replace_in_runs(p, old, new_parts):
    """Replace `old` (inside one run) by a sequence of ('text'|('tok', NAME))."""
    for i, r in enumerate(p.runs):
        if old in r.text:
            before, after = r.text.split(old, 1)
            r.text = before
            anchor = r._r
            for part in new_parts:
                nr = copy.deepcopy(anchor)
                anchor.addnext(nr)
                anchor = nr
                nrun = docx.text.run.Run(nr, p)
                if isinstance(part, tuple):
                    nrun.text = tok(part[1])
                    _style_token(nrun)
                else:
                    nrun.text = part
            nr = copy.deepcopy(anchor)
            anchor.addnext(nr)
            docx.text.run.Run(nr, p).text = after
            return True
    raise ValueError(f"{old!r} not found in one run of: {p.text[:80]!r}")


def append_token(p, name, prefix=" "):
    r = p.add_run(prefix + tok(name))
    _style_token(r)
    return r


def prepend_token(p, name):
    first = p.runs[0]
    nr = copy.deepcopy(first._r)
    first._r.addprevious(nr)
    run = docx.text.run.Run(nr, p)
    run.text = tok(name) + " "
    _style_token(run)


def set_text(p, parts):
    """Rewrite a simple paragraph as text + tokens, keeping the first run's style."""
    base = p.runs[0]
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    base.text = ""
    anchor = base._r
    first = True
    for part in parts:
        if first:
            run = base
            first = False
        else:
            nr = copy.deepcopy(base._r)
            anchor.addnext(nr)
            anchor = nr
            run = docx.text.run.Run(nr, p)
        if isinstance(part, tuple):
            run.text = tok(part[1])
            _style_token(run)
        else:
            run.text = part
            run.font.color.rgb = None
            run.font.size = None


T = lambda n: ("tok", n)


def do_tr():
    d = docx.Document("TR.docx")
    t = d.tables[0]
    c = lambda r, col, i: t.rows[r].cells[col].paragraphs[i]
    replace_in_runs(c(1, 0, 1), "S25328", [T("FILENO")])
    replace_in_runs(c(1, 1, 0), "S25328", [T("FILENO_FA")])
    replace_in_runs(c(2, 0, 0), "31/10/2025", [T("DATE")])
    replace_in_runs(c(2, 0, 0), ", located at ", [" ", T("CLIENT"), ", located at ", T("ADDR")])
    replace_in_runs(c(2, 1, 0), "31/10/2025", [T("DATE_FA")])
    p = c(2, 1, 0)
    # "و متقاضی،  که به نشانی: " -> insert name after متقاضی، and address at the end
    replace_in_runs(p, "،  ", ["، ", T("CLIENT_FA"), " "])
    append_token(p, "ADDR_FA", prefix=" ")
    for i, name in ((4, "CB_SP"), (5, "CB_WP"), (6, "CB_TRV")):
        prepend_token(c(5, 0, i), name)
        prepend_token(c(5, 1, i), name + "_FA")
    for i, name in ((5, "PROFEE"), (6, "GOVFEE"), (7, "BIOFEE")):
        replace_in_runs(c(9, 0, i), "Fee  CAD", [T(name), "  CAD"])
    replace_in_runs(c(9, 0, 8), "Fee CAD", [T("TAX"), " CAD"])
    replace_in_runs(c(9, 0, 9), "Fees  CAD", [T("TOTAL"), "  CAD"])
    replace_in_runs(c(9, 1, 5), "Fee", [T("PROFEE_FA")])
    replace_in_runs(c(9, 1, 6), "Fee", [T("GOVFEE_FA")])
    p = c(9, 1, 7)
    replace_in_runs(p, "Fees", [T("TOTAL_FA")])
    replace_in_runs(p, "Fee", [T("BIOFEE_FA")])
    replace_in_runs(p, "Fee", [T("TAX_FA")])
    replace_in_runs(c(11, 0, 2), "...........", [T("PAY1")])
    replace_in_runs(c(11, 0, 3), "...........", [T("PAY2")])
    set_text(c(11, 0, 4), [T("PAYNOTE")])
    p = c(11, 1, 1)
    replace_in_runs(p, "...........", [T("PAY1_FA")])
    replace_in_runs(p, "...........", [T("PAY2_FA")])
    for i, name in ((2, "GIVEN"), (3, "FAMILY"), (4, "ADDR2"), (5, "PHONE"), (6, "EMAIL")):
        append_token(c(24, 0, i), name)
        append_token(c(24, 1, i), name + "_FA")
    set_text(c(26, 0, 3), [T("SIG_C"), "\t\t\t\t", T("SIG_R")])
    p = c(26, 0, 5)
    set_text(p, ["Signature of Client \t\tDate: ", T("DATE_C"), "\t\t\tSignature of RCIC           Date: ", T("DATE_R")])
    d.save("TR_t.docx")


def do_pr():
    d = docx.Document("PR.docx")
    P = d.paragraphs
    append_token(P[3], "FILENO", prefix="")
    replace_in_runs(P[5], "Date", [T("DATE")])
    set_text(P[6], [", and the Client ", T("CLIENT"), ", located at ", T("ADDR")])
    for i, name in ((16, "CB_EE"), (17, "CB_PNP"), (18, "CB_PNPEE")):
        prepend_token(P[i], name)
    replace_in_runs(P[52], "$Fee", ["$", T("PROFEE")])
    replace_in_runs(P[54], "$Fee", ["$", T("TAX")])
    replace_in_runs(P[56], "$Amount", ["$", T("DISCOUNT")])
    replace_in_runs(P[58], "$Fee", ["$", T("TOTAL")])
    replace_in_runs(P[66], "$Fee", ["$", T("PAY1")])
    replace_in_runs(P[67], "$Fee", ["$", T("PAY2")])
    set_text(P[68], [T("PAYNOTE")])
    for i, name in ((129, "GIVEN"), (130, "FAMILY"), (131, "ADDR2"), (132, "PHONE"), (133, "EMAIL")):
        append_token(P[i], name)
    set_text(P[148], [T("SIG_C"), "\t\t\t\t\t\t", T("SIG_R")])
    set_text(P[150], ["Signature of Client \t\t\t\t\t\tSignature of RCIC    Date ", T("DATE_R")])
    set_text(P[151], ["Date ", T("DATE_C")])
    d.save("PR_t.docx")


if __name__ == "__main__":
    do_tr()
    do_pr()
    print("tokenized: TR_t.docx PR_t.docx")
